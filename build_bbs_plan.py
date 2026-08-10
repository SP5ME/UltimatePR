from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.enum.section import WD_ORIENT
from pathlib import Path

OUT = Path(r"H:\BBS Serwer\Plan_projektu_nowoczesnego_BBS_Packet_Radio.docx")

NAVY = "16324F"
BLUE = "2474A6"
CYAN = "16A3B6"
LIGHT = "EAF3F7"
PALE = "F5F8FA"
MID = "D5E1E8"
GRAY = "5B6770"
DARK = "17212B"
WHITE = "FFFFFF"
GREEN = "287D5B"
AMBER = "B87816"


def font(run, size=10, bold=False, color=DARK, italic=False, name="Aptos"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=120, start=140, bottom=120, end=140):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent=120):
    total = sum(widths_dxa)
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(total)); tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd"); tblPr.append(tblInd)
    tblInd.set(qn("w:w"), str(indent)); tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(width)); grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW"); tcPr.append(tcW)
            tcW.set(qn("w:w"), str(widths_dxa[idx])); tcW.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            margins(cell)


def repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def keep_with_next(p):
    p.paragraph_format.keep_with_next = True


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text); font(r, 9.5)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); font(r, 9.5)
    return p


def add_label_para(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.12
    a = p.add_run(label + " "); font(a, 10, True, NAVY)
    b = p.add_run(text); font(b, 10)
    return p


def callout(doc, label, text, color=BLUE):
    t = doc.add_table(rows=1, cols=1)
    set_table_geometry(t, [9360], 120)
    c = t.cell(0, 0); shade(c, LIGHT); margins(c, 180, 220, 180, 220)
    p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label.upper() + "\n"); font(r, 8.5, True, color)
    r = p.add_run(text); font(r, 10.5, True, NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def notes_block(doc, lines=7):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7); p.paragraph_format.space_after = Pt(4)
    r = p.add_run("UWAGI / DECYZJE / PYTANIA"); font(r, 9, True, BLUE)
    keep_with_next(p)
    t = doc.add_table(rows=lines, cols=1)
    set_table_geometry(t, [9360], 120)
    for row in t.rows:
        c = row.cells[0]; margins(c, 80, 140, 80, 140)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        c.paragraphs[0].paragraph_format.space_after = Pt(0)
        c.paragraphs[0].add_run(" ")


stages = [
    {
        "n":"0", "title":"Założenia, zgodność i clean-room", "milestone":"Zatwierdzona specyfikacja produktu i zasady niezależnej implementacji",
        "goal":"Ustalić zakres pierwszego wydania, źródła specyfikacji i granice prawne. Projekt nie kopiuje kodu LinBPQ ani URONode.",
        "scope":["Katalog wymagań funkcjonalnych node'a i BBS", "Rejestr publicznych specyfikacji: AX.25, KISS, NET/ROM i FBB", "Zasady clean-room oraz ewidencja pochodzenia wiedzy", "Wybór licencji nowego projektu", "Słownik pojęć Packet Radio i model kompatybilności"],
        "accept":["Każda funkcja MVP ma opis wejść, wyjść i oczekiwanego zachowania", "Źródła wiedzy są zapisane i możliwe do audytu", "Zakres pierwszej wersji jest zamknięty", "Nie ma skopiowanych struktur ani fragmentów kodu z analizowanych projektów"],
        "tests":["Przegląd wymagań przez krótkofalowca/operatora BBS", "Macierz funkcja–specyfikacja–test", "Kontrola licencji wszystkich zależności"],
        "risks":"Największe ryzyko to niekontrolowane rozszerzanie zakresu oraz odtwarzanie zachowań bez jednoznacznej specyfikacji."
    },
    {
        "n":"1", "title":"Fundament projektu i panel WWW", "milestone":"Serwer uruchamia się na Windows i Linux, a konfiguracja działa przez przeglądarkę",
        "goal":"Zbudować szkielet w Go, stabilny model konfiguracji, API oraz pierwszy panel administracyjny.",
        "scope":["Struktura modułów i repozytorium", "Konfiguracja w SQLite z migracjami", "Logowanie administratora i sesje", "Dashboard stanu procesu", "Rejestr zdarzeń i logów", "Budowanie dla Windows, Linux x86-64 i ARM64"],
        "accept":["Pierwsze uruchomienie prowadzi przez kreator stacji", "Konfiguracja przetrwa restart", "Panel nie wymaga dostępu do Internetu", "Powstają powtarzalne binaria dla obsługiwanych platform"],
        "tests":["Testy API i migracji bazy", "Test logowania, wylogowania i blokady prób", "Uruchomienie na Windows i Debian VM", "Test przywrócenia kopii konfiguracji"],
        "risks":"Zbyt wczesne rozbudowanie interfejsu może odciągnąć pracę od warstwy radiowej."
    },
    {
        "n":"2", "title":"KISS i monitor ramek AX.25", "milestone":"Pierwsze ramki z Direwolf/TNC są widoczne w panelu",
        "goal":"Dostarczyć niezależną warstwę transportową oraz bezpieczny dekoder i enkoder ramek AX.25.",
        "scope":["KISS TCP", "KISS przez port szeregowy", "Framing, escaping i obsługa błędnych ramek", "Kodowanie/dekodowanie adresów AX.25 i SSID", "Monitor RX/TX w czasie rzeczywistym", "Pliki przechwyconych ramek do odtwarzania testów"],
        "accept":["Połączenie z Direwolf działa stabilnie", "Ramki UI i sterujące są poprawnie prezentowane", "Uszkodzona ramka nie zatrzymuje serwera", "Można filtrować ruch według portu i znaku"],
        "tests":["Golden vectors ramek AX.25", "Fuzzing dekodera", "Rozłączanie i ponowne łączenie KISS TCP", "Test długotrwały z zapisem RX/TX"],
        "risks":"Błędy parsera ramek mogą prowadzić do niezgodności lub awarii przy nietypowym ruchu radiowym."
    },
    {
        "n":"3", "title":"Połączeniowy AX.25", "milestone":"Dwie instancje zestawiają stabilną sesję AX.25",
        "goal":"Zaimplementować maszynę stanów połączenia AX.25 wraz z retransmisją i kontrolą przepływu.",
        "scope":["SABM/SABME, UA, DISC i DM", "Ramki I, RR, RNR i REJ", "Numerowanie ramek i okno transmisji", "Timery T1, T2, T3 oraz limit N2", "Segmentacja danych zgodnie z PACLEN", "Wiele równoczesnych sesji"],
        "accept":["Połączenie, transmisja dwukierunkowa i rozłączenie są powtarzalne", "Utrata ramek uruchamia poprawną retransmisję", "Stan sesji jest widoczny w panelu", "Restart nie pozostawia zablokowanych zasobów"],
        "tests":["Symulator utraty, opóźnienia, duplikacji i zmiany kolejności", "Testy maszyn stanów", "Interoperacyjność z LinBPQ i innym TNC", "Test wielu jednoczesnych sesji"],
        "risks":"To najbardziej krytyczny technicznie etap; drobne błędy czasowe mogą ujawniać się tylko na wolnym lub zakłóconym łączu."
    },
    {
        "n":"4", "title":"Node Packet Radio", "milestone":"Użytkownik radiowy może wejść do node'a i zestawić dalsze połączenie",
        "goal":"Zbudować czytelną powłokę node'a, routing między portami i podstawowe narzędzia operatorskie.",
        "scope":["Prompt, alias i komunikaty stacji", "Komendy CONNECT, PORTS, USERS, MHEARD, INFO, HELP i BYE", "Rejestr komend i uprawnień", "Połączenia wychodzące przez wskazany port", "Routing sesji między portami", "Konfigurowalne limity i timeouty"],
        "accept":["Komendy działają zarówno przez radio, jak i lokalny terminal testowy", "Błędne polecenia nie destabilizują sesji", "MHEARD i aktywne sesje są widoczne w WWW", "Operator może bezpiecznie zakończyć sesję"],
        "tests":["Testy parsera poleceń", "Scenariusze uprawnień", "Połączenie wieloskokowe w laboratorium", "Test zgodności tekstowego interfejsu z terminalami packet"],
        "risks":"Rozbudowane aliasy i komendy zewnętrzne mogą stać się źródłem podatności, jeśli otrzymają dostęp do powłoki systemowej."
    },
    {
        "n":"5", "title":"BBS i magazyn wiadomości", "milestone":"Użytkownik może wysłać, wylistować, odczytać i usunąć wiadomość",
        "goal":"Dostarczyć podstawowy, niezawodny BBS z prywatnymi wiadomościami, biuletynami i panelem zarządzania.",
        "scope":["Użytkownicy, znaki, role i Home BBS", "Wiadomości prywatne i biuletyny", "Komendy L, R, S, K, B oraz HELP", "Nagłówki, identyfikatory BID/MID i statusy", "Wyszukiwanie i administracja przez WWW", "Backup, eksport i kontrola integralności SQLite"],
        "accept":["Pełny cykl wiadomości działa przez AX.25", "Duplikaty BID są wykrywane", "Treść pozostaje poprawna po restarcie", "Administrator może wykonać i odtworzyć backup"],
        "tests":["Scenariusze prywatne i bulletin", "Duże wiadomości i nietypowe zakończenia linii", "Awaria podczas zapisu", "Migracje schematu i odtwarzanie kopii"],
        "risks":"Model danych musi uwzględniać późniejszy forwarding; zbyt prosty schemat wymusi kosztowne migracje."
    },
    {
        "n":"6", "title":"Forwarding BBS/FBB", "milestone":"Wiadomości są automatycznie wymieniane z istniejącym BBS-em",
        "goal":"Zapewnić interoperacyjny forwarding, kolejki, routing pocztowy i pełną obserwowalność wymiany.",
        "scope":["SID i negocjacja możliwości", "Klasyczny forwarding tekstowy", "Propozycje FBB", "Forwarding blokowy i kompresja", "Partnerzy, harmonogramy i skrypty połączeń", "Reguły routingu wiadomości P/B/T", "Kolejki, retry, raporty i ręczne uruchamianie"],
        "accept":["Wymiana z LinBPQ działa w obu kierunkach", "Brak duplikatów i pętli routingu", "Przerwana wymiana jest bezpiecznie wznawiana", "Panel pokazuje powód przyjęcia, odrzucenia lub oczekiwania"],
        "tests":["Macierz interoperacyjności z kilkoma BBS-ami", "Przerwanie sesji w każdym kroku", "Duplikaty, odrzucenia i limity rozmiaru", "Testy kompresji na znanych wektorach"],
        "risks":"Różne implementacje FBB mają historyczne rozszerzenia i odstępstwa; potrzebne będą testy rzeczywiste, nie tylko specyfikacja."
    },
    {
        "n":"7", "title":"NET/ROM i funkcje sieciowe", "milestone":"Node uczestniczy w sieci wielowęzłowej i wybiera trasę",
        "goal":"Rozszerzyć serwer o funkcje sieciowe potrzebne do pracy w większej sieci Packet Radio.",
        "scope":["NET/ROM i beacony węzłów", "Tablice sąsiadów, tras i jakości", "Routing statyczny i dynamiczny", "Digipeater i zasady ścieżek", "Opcjonalne interfejsy Linux AX.25", "Późniejsze rozszerzenia: ROSE, APRS, TCP services i Winlink"],
        "accept":["Trasy są stabilne i możliwe do wyjaśnienia w panelu", "Zmiany sąsiadów nie powodują pętli", "Operator może ustawić ograniczenia i preferencje", "Funkcje dodatkowe można wyłączać modułowo"],
        "tests":["Topologie 3–5 węzłów w symulatorze", "Zanik i powrót łącza", "Pętle oraz konflikt aliasów", "Testy interoperacyjności w wydzielonym środowisku"],
        "risks":"Nie należy dopuścić, aby rozszerzenia sieciowe opóźniły stabilne wydanie podstawowego BBS."
    },
    {
        "n":"8", "title":"Utwardzenie, pakowanie i wydanie", "milestone":"Instalowalne wydanie dla Debian/Raspberry Pi OS oraz zweryfikowany wariant Alpine",
        "goal":"Przygotować produkt do długotrwałej pracy bez nadzoru i łatwej instalacji przez krótkofalowca.",
        "scope":["Pakiety .deb i usługa systemd", "Wariant Alpine/OpenRC po walidacji musl", "Aktualizacje, migracje i rollback", "Diagnostyka do pobrania z panelu", "TLS, role, audyt i bezpieczne domyślne ustawienia", "Dokumentacja operatora oraz obrazy laboratoryjne"],
        "accept":["Instalacja na czystej maszynie wymaga minimalnej liczby kroków", "Aktualizacja zachowuje wiadomości i konfigurację", "Po awarii serwer automatycznie wraca do pracy", "Testy 24/7 i kontrola bezpieczeństwa kończą się bez krytycznych błędów"],
        "tests":["Debian x86-64, Raspberry Pi ARM64 i Alpine", "Długotrwały soak test", "Backup/restore po aktualizacji", "Skan zależności, test uprawnień i próby błędnej konfiguracji"],
        "risks":"Pakowanie sprzętowych interfejsów USB, portów szeregowych i capabilities wymaga osobnych profili dla różnych instalacji."
    },
]


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5); sec.page_height = Inches(11)
sec.top_margin = Inches(0.78); sec.bottom_margin = Inches(0.72)
sec.left_margin = Inches(1); sec.right_margin = Inches(1)
sec.header_distance = Inches(0.32); sec.footer_distance = Inches(0.32)

# Styles: compact_reference_guide resolved tokens
styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Aptos"; normal.font.size = Pt(10); normal.font.color.rgb = RGBColor.from_string(DARK)
normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos"); normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
normal.paragraph_format.space_after = Pt(5); normal.paragraph_format.line_spacing = 1.12
for name, size, color, before, after in [("Title",30,NAVY,0,8),("Subtitle",13,GRAY,0,10),("Heading 1",19,NAVY,12,6),("Heading 2",12,BLUE,9,4),("Heading 3",10,CYAN,6,3)]:
    st=styles[name]; st.font.name="Aptos Display" if name in ("Title","Heading 1") else "Aptos"; st.font.size=Pt(size); st.font.bold=name!="Subtitle"; st.font.color.rgb=RGBColor.from_string(color)
    st._element.rPr.rFonts.set(qn("w:ascii"), st.font.name); st._element.rPr.rFonts.set(qn("w:hAnsi"), st.font.name)
    st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True

# Header/footer
hp = sec.header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
r=hp.add_run("NOWOCZESNY BBS PACKET RADIO"); font(r,8,True,BLUE)
fp = sec.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=fp.add_run("Plan projektu  •  wersja robocza 1.0  •  sierpień 2026"); font(r,8,False,GRAY)

# Cover - editorial_cover
doc.add_paragraph().paragraph_format.space_after=Pt(54)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("PLAN PROJEKTU"); font(r,10,True,CYAN)
p=doc.add_paragraph(style="Title"); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run("Nowoczesny serwer BBS\nPacket Radio")
p=doc.add_paragraph(style="Subtitle"); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Niezależna implementacja inspirowana doświadczeniami LinBPQ i URONode")
doc.add_paragraph().paragraph_format.space_after=Pt(14)
callout(doc,"Kierunek projektu","Nowy kod od podstaw • Go • panel WWW • AX.25/KISS • Linux, Raspberry Pi i Alpine")
doc.add_paragraph().paragraph_format.space_after=Pt(14)
t=doc.add_table(rows=4,cols=2); set_table_geometry(t,[2200,7160],120)
meta=[("Dokument","Plan pracy, kamienie milowe i miejsce na uwagi"),("Zakres","Node Packet Radio, BBS, forwarding i zarządzanie WWW"),("Platformy","Windows (development), Debian/Raspberry Pi OS, Alpine"),("Status","Dokument roboczy do wspólnego rozwijania")]
for i,(a,b) in enumerate(meta):
    shade(t.cell(i,0),PALE); margins(t.cell(i,0),120,170,120,170); margins(t.cell(i,1),120,170,120,170)
    p=t.cell(i,0).paragraphs[0]; p.paragraph_format.space_after=Pt(0); font(p.add_run(a.upper()),8.5,True,BLUE)
    p=t.cell(i,1).paragraphs[0]; p.paragraph_format.space_after=Pt(0); font(p.add_run(b),9.5)

doc.add_page_break()
doc.add_heading("1. Krótkie podsumowanie etapów", level=1)
callout(doc,"Strategia","Najpierw sprawdzalny tor radiowy KISS/AX.25, następnie node i BBS, potem forwarding oraz funkcje sieciowe. Panel WWW rozwijamy od początku jako narzędzie konfiguracji i diagnostyki.",GREEN)

summary=doc.add_table(rows=1,cols=4); summary.alignment=WD_TABLE_ALIGNMENT.LEFT
set_table_geometry(summary,[780,2900,3540,2140],120)
headers=["Etap","Obszar","Najważniejszy rezultat","Status"]
for i,h in enumerate(headers):
    c=summary.rows[0].cells[i]; shade(c,NAVY); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(0); font(p.add_run(h),8.5,True,WHITE)
repeat_header(summary.rows[0])
for s in stages:
    cells=summary.add_row().cells
    vals=[s["n"],s["title"],s["milestone"],"☐ Plan\n☐ Praca\n☐ Gotowe"]
    for i,v in enumerate(vals):
        if int(s["n"])%2: shade(cells[i],PALE)
        cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p=cells[i].paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.alignment=WD_ALIGN_PARAGRAPH.CENTER if i in (0,3) else WD_ALIGN_PARAGRAPH.LEFT
        font(p.add_run(v),8.4, i==0, NAVY if i==0 else DARK)
set_table_geometry(summary,[780,2900,3540,2140],120)

doc.add_heading("Zasady prowadzenia projektu", level=2)
for x in ["Każdy etap kończy się działającym, możliwym do demonstracji przyrostem.","Najpierw testy protokołu i zgodności, potem rozbudowa interfejsu.","Kod powstaje niezależnie; LinBPQ i URONode służą do poznania funkcji i testów interoperacyjności, nie jako źródło kodu.","Funkcje opcjonalne nie mogą destabilizować podstawowego node'a i BBS.","Dokument należy aktualizować po każdej istotnej decyzji technicznej."]:
    add_bullet(doc,x)

doc.add_heading("Środowisko testowe", level=2)
add_label_para(doc,"Windows:","panel, API, testy automatyczne, Direwolf oraz dwie lokalne instancje serwera.")
add_label_para(doc,"Debian VM:","kompilacja i test integracyjny usług systemowych, portów oraz długiej pracy.")
add_label_para(doc,"Radio/TNC:","końcowe testy interoperacyjności przez KISS TCP lub przekazany port USB/Serial.")

doc.add_page_break()
doc.add_heading("2. Organizacja i kryteria jakości", level=1)
doc.add_heading("Proponowane moduły", level=2)
modules=[("transport","KISS TCP/Serial, później Linux AX.25 i inne transporty"),("ax25","ramki, adresy, maszyna stanów i timery"),("session","sesje użytkowników oraz kontrola przepływu"),("node","komendy, routing połączeń i MHEARD"),("bbs","użytkownicy, wiadomości, BID/MID i wyszukiwanie"),("forwarding","FBB, kolejki, harmonogram i reguły trasowania"),("web","panel administracyjny, API i monitoring"),("storage","SQLite, migracje, backup i integralność")]
t=doc.add_table(rows=1,cols=2); set_table_geometry(t,[2300,7060],120)
for i,h in enumerate(["Moduł","Odpowiedzialność"]): shade(t.cell(0,i),NAVY); p=t.cell(0,i).paragraphs[0]; p.paragraph_format.space_after=Pt(0); font(p.add_run(h),9,True,WHITE)
repeat_header(t.rows[0])
for idx,(a,b) in enumerate(modules):
    c=t.add_row().cells
    if idx%2: shade(c[0],PALE); shade(c[1],PALE)
    p=c[0].paragraphs[0]; p.paragraph_format.space_after=Pt(0); font(p.add_run(a),9,True,BLUE)
    p=c[1].paragraphs[0]; p.paragraph_format.space_after=Pt(0); font(p.add_run(b),9)
set_table_geometry(t,[2300,7060],120)

doc.add_heading("Definicja ukończenia etapu", level=2)
for x in ["Kod przechodzi testy automatyczne i kontrolę statyczną.","Funkcja ma scenariusz demonstracyjny i opis konfiguracji.","Błędy wejściowe nie powodują awarii ani utraty danych.","Zmiana działa na Windows (jeśli dotyczy) oraz docelowym Linuxie.","Zapisano decyzje, znane ograniczenia i zadania przeniesione dalej."]:
    add_bullet(doc,x)

doc.add_heading("Rejestr decyzji ogólnych", level=2)
notes_block(doc,9)

# Stage worksheets
for s in stages:
    doc.add_page_break()
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2)
    r=p.add_run(f"ETAP {s['n']}"); font(r,9,True,CYAN)
    p=doc.add_paragraph(style="Heading 1"); p.paragraph_format.space_before=Pt(0); p.add_run(s["title"])
    callout(doc,"Kamień milowy",s["milestone"],GREEN)
    add_label_para(doc,"Cel:",s["goal"])

    doc.add_heading("Zakres prac", level=2)
    for x in s["scope"]: add_bullet(doc,x)

    doc.add_heading("Kryteria odbioru", level=2)
    for x in s["accept"]: add_bullet(doc,"☐ " + x)

    doc.add_heading("Plan testów", level=2)
    for x in s["tests"]: add_bullet(doc,"☐ " + x)

    doc.add_heading("Ryzyko etapu", level=2)
    add_label_para(doc,"Uwaga:",s["risks"])

    # tracking strip
    t=doc.add_table(rows=2,cols=4); set_table_geometry(t,[2340,2340,2340,2340],120)
    vals=[("Właściciel",""),("Planowana data",""),("Data zakończenia",""),("Status","☐ Plan  ☐ Praca  ☐ Gotowe")]
    for i,(a,b) in enumerate(vals):
        shade(t.cell(0,i),MID); shade(t.cell(1,i),PALE)
        p=t.cell(0,i).paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(0); font(p.add_run(a.upper()),8,True,NAVY)
        p=t.cell(1,i).paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(0); font(p.add_run(b if b else " "),8.5)
    notes_block(doc,8)

doc.add_page_break()
doc.add_heading("3. Backlog funkcji po pierwszym wydaniu", level=1)
callout(doc,"Zasada","Poniższe elementy trafiają do pierwszej wersji tylko wtedy, gdy nie opóźniają stabilnego KISS/AX.25, node'a, BBS i forwardingu.",AMBER)
for title,items in [
    ("Protokoły i sieć",["ROSE i FlexNet","APRS digipeater oraz APRS-IS","IP over AX.25 / AMPRNet","Interfejs AGWPE","MQTT i zewnętrzne API zdarzeń"]),
    ("Poczta i usługi",["WebMail użytkownika","SMTP/POP3/NNTP","Winlink CMS i format B2F","Formularze wiadomości","Zaawansowane raporty ruchu"]),
    ("Eksploatacja",["Aktualizacje z panelu","Obrazy gotowe dla Raspberry Pi","Kontenery z profilami urządzeń","Klaster/HA dla zastosowań eksperymentalnych","Wtyczki rozszerzające typy portów"])
]:
    doc.add_heading(title,level=2)
    for x in items: add_bullet(doc,"☐ " + x)

doc.add_heading("Uwagi do backlogu", level=2)
notes_block(doc,10)

doc.add_page_break()
doc.add_heading("4. Dziennik ustaleń", level=1)
p=doc.add_paragraph("Miejsce na ważne decyzje podjęte w trakcie realizacji. Każdy wpis powinien zawierać datę, decyzję, uzasadnienie i wpływ na dalsze etapy.")
p.paragraph_format.space_after=Pt(8)
t=doc.add_table(rows=1,cols=4); set_table_geometry(t,[1200,2500,3620,2040],120)
for i,h in enumerate(["Data","Decyzja","Uzasadnienie","Wpływ / autor"]): shade(t.cell(0,i),NAVY); p=t.cell(0,i).paragraphs[0]; p.paragraph_format.space_after=Pt(0); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(p.add_run(h),8.5,True,WHITE)
repeat_header(t.rows[0])
for _ in range(13):
    row=t.add_row()
    for c in row.cells:
        margins(c,170,140,170,140); c.paragraphs[0].add_run(" ")
set_table_geometry(t,[1200,2500,3620,2040],120)

# Core properties
doc.core_properties.title = "Plan projektu nowoczesnego serwera BBS Packet Radio"
doc.core_properties.subject = "Plan pracy, etapy, testy i miejsce na uwagi"
doc.core_properties.author = "Zespół projektu BBS Packet Radio"
doc.core_properties.keywords = "Packet Radio, BBS, AX.25, KISS, FBB, Go"

doc.save(OUT)
print(OUT)
